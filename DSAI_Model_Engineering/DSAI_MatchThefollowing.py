########################################################

#Copyright (c) DeepSphere.AI 2021

# All rights reserved

# We are sharing this partial code for learning and research, and the idea behind us sharing the source code is to stimulate ideas #and thoughts for the learners to develop their MLOps.

# Author: # DeepSphere.AI | deepsphere.ai | dsschoolofai.com | info@deepsphere.ai

# Release: Initial release

#######################################################


import json
import requests
import string
import re
import nltk
import string
import itertools
import streamlit as st
# nltk.download('stopwords')
# nltk.download('wordnet')
# nltk.download('punkt')
import pke
from nltk.corpus import stopwords
from nltk.corpus import wordnet
import traceback
from flashtext import KeywordProcessor
import re
from pprint import pprint
import random
import pandas as pd
from fuzzywuzzy import fuzz
from docxtpl import DocxTemplate
import docx
# from prettytable import PrettyTable
# from IPython.display import Markdown, display

def file_selector_match():
    vAR_file = st.file_uploader('Upload the text file',type=['txt'],key='3')
    if vAR_file is not None:
        vAR_text = vAR_file.read().decode("utf-8")
        st.write('File Content: '+ vAR_text)
        return vAR_text

def sentence_answers(keyword_sentence_mapping):
    vAR_solution = {}
    for k,v in keyword_sentence_mapping.items():
        if v:
            vAR_match = v[0].lower()
            if k in vAR_match:
                vAR_temp = re.compile(re.escape(k), re.IGNORECASE)
                vAR_solution[k] = vAR_temp.sub('<answer>',vAR_match)
            else:
                vAR_solution[k] = vAR_match
    return vAR_solution

def quest_ans(sol):
    sol_df = pd.DataFrame(list(sol.items()),columns= ["Column A", "Column B"])
    sol_df["Column B"] = sol_df["Column B"].apply(lambda x: re.sub(r"\n+"," ",x))

    sh_val = list(sol.values())
    random.shuffle(sh_val)
    new_dic = dict(zip(sol,sh_val))
    quest_df = pd.DataFrame(list(new_dic.items()),columns= ["Column A", "Column B"])
    quest_df["Column B"] = quest_df["Column B"].apply(lambda x: re.sub(r"\n+"," ",x))

    return quest_df, sol_df

@st.cache(show_spinner=False)
# def get_keywords(text):
#     vAR_out=[]
#     try:
#         vAR_extractor = pke.unsupervised.YAKE()
#         vAR_extractor.load_document(input=text)
#         # pos = {'VERB', 'ADJ', 'NOUN'}
#         vAR_pos ={'NOUN'}
#         vAR_stoplist = list(string.punctuation)
#         vAR_stoplist += ['-lrb-', '-rrb-', '-lcb-', '-rcb-', '-lsb-', '-rsb-']
#         vAR_stoplist += stopwords.words('english')
#         vAR_extractor.candidate_selection(n=2,pos=vAR_pos, stoplist=vAR_stoplist)

#         vAR_extractor.candidate_weighting(window=3,
#                                       stoplist=vAR_stoplist,
#                                       use_stems=False)

#         vAR_keyphrases = vAR_extractor.get_n_best(n=30)
        
#         for val in vAR_keyphrases:
#             vAR_out.append(val[0])
#     except:
#         vAR_out = []
#         traceback.print_exc()

#     return vAR_out

def get_noun_adj_verb_sr(text):
  # define the set of valid Part-of-Speeches
  pos = {'NOUN', 'PROPN', 'ADJ'}

  # 1. create a SingleRank extractor.
  extractor = pke.unsupervised.SingleRank()

  # 2. load the content of the document.
  extractor.load_document(input=text,
                          language='en',
                          normalization=None)

  # 3. select the longest sequences of nouns and adjectives as candidates.
  extractor.candidate_selection(pos=pos)

  # 4. weight the candidates using the sum of their word's scores that are
  #    computed using random walk. In the graph, nodes are words of
  #    certain part-of-speech (nouns and adjectives) that are connected if
  #    they occur in a window of 10 words.
  extractor.candidate_weighting(window=10,
                                pos=pos)

  # 5. get the 10-highest scored candidates as keyphrases
  keyphrases = extractor.get_n_best(n=50)
  keyphrases_res = zip(*keyphrases)
  keyphrases = list(list(keyphrases_res)[0])
  z = []
  keyphrases.sort()
  for x, y in zip(keyphrases, keyphrases[1:]):
    if fuzz.partial_ratio(x,y)!=100:
        z.append(x)
  return z


def get_keywords_mtf(text):
      # 1. create a WINGNUS extractor.
  extractor = pke.supervised.WINGNUS()

  # 2. load the content of the document.
  extractor.load_document(input=text)

  # 3. select simplex noun phrases as candidates.
  extractor.candidate_selection()

  # 4. classify candidates as keyphrase or not keyphrase.
  df = pke.load_document_frequency_file(input_file='df-semeval2010.tsv.gz')
  model_file = 'WINGNUS-semeval2010.py3.pickle'
  extractor.candidate_weighting( model_file=model_file, df=df)

  # 5. get the 10-highest scored candidates as keyphrases
  keyphrases = extractor.get_n_best(n=50)
  wingnus_res = zip(*keyphrases)
  keyphrases = list(list(wingnus_res)[0])
  z = []
  keyphrases.sort()
  print(keyphrases)
  for x, y in zip(keyphrases, keyphrases[1:]):
    print(x,y)
    if fuzz.partial_ratio(x,y)!=100:
        z.append(x)
  z = set(z)
  sr = set(get_noun_adj_verb_sr(text))
  common = sr.intersection(z)
  print('-----------------Keywords from Single Rank in MTF---------------------\n',sr)
  print('-----------------Keywords from wingnus in MTF---------------------\n',z)
  print('-----------------Common Keywords from both in MTF---------------------\n',common)
  return common

#       # define the set of valid Part-of-Speeches
#   pos = {'NOUN', 'PROPN', 'ADJ'}

#   # 1. create a SingleRank extractor.
#   extractor = pke.unsupervised.SingleRank()

#   # 2. load the content of the document.
#   extractor.load_document(input=text,
#                           language='en',
#                           normalization=None)

#   # 3. select the longest sequences of nouns and adjectives as candidates.
#   extractor.candidate_selection(pos=pos)

#   # 4. weight the candidates using the sum of their word's scores that are
#   #    computed using random walk. In the graph, nodes are words of
#   #    certain part-of-speech (nouns and adjectives) that are connected if
#   #    they occur in a window of 10 words.
#   extractor.candidate_weighting(window=10,
#                                 pos=pos)

#   # 5. get the 10-highest scored candidates as keyphrases
#   keyphrases = extractor.get_n_best(n=20)
#   keyphrases_res = zip(*keyphrases)
#   keyphrases = list(list(keyphrases_res)[0])
#   z = []
#   keyphrases.sort()
# #   print(keyphrases)
#   for x, y in zip(keyphrases, keyphrases[1:]):
#     print(x, y)
#     if fuzz.partial_ratio(x,y)!=100:
#       z.append(x)
#   return z
  
#Extract sentences having the keywords that is extracted before.
@st.cache(show_spinner=False)
def get_sentences_for_keyword(keywords, sentences):
    vAR_keyword_processor = KeywordProcessor()
    vAR_keyword_sentences = {}
    for word in keywords:
        vAR_keyword_sentences[word] = []
        vAR_keyword_processor.add_keyword(word)
    for sentence in sentences:
        vAR_keywords_found = vAR_keyword_processor.extract_keywords(sentence)
        for key in vAR_keywords_found:
            vAR_keyword_sentences[key].append(sentence)

    for key in vAR_keyword_sentences.keys():
        vAR_values = vAR_keyword_sentences[key]
        vAR_values = sorted(vAR_values, key=len, reverse=False)
        vAR_keyword_sentences[key] = vAR_values
    return vAR_keyword_sentences
def write_table(path,quest_df,sol_df):
  doc = docx.Document(path)

  t_q = doc.add_table(quest_df.shape[0]+1, quest_df.shape[1])
  for j in range(quest_df.shape[-1]):
      t_q.cell(0,j).text = quest_df.columns[j]

  for i in range(quest_df.shape[0]):
      for j in range(quest_df.shape[-1]):
          t_q.cell(i+1,j).text = str(quest_df.values[i,j])

  doc.add_paragraph("\nANSWER")
  t_a = doc.add_table(sol_df.shape[0]+1, sol_df.shape[1])
  for j in range(sol_df.shape[-1]):
      t_a.cell(0,j).text = sol_df.columns[j]

  for i in range(sol_df.shape[0]):
      for j in range(sol_df.shape[-1]):
          t_a.cell(i+1,j).text = str(sol_df.values[i,j])

  doc.save(path)
def get_context_mtf(unit,title,grade):
  return {
      "unit": unit,
      # "grade" : grade+"th Grade",
      "grade" : grade+"th Grade",
      "title" : title.title(),
  }

# def sentence_answers(keyword_sentence_mapping):
#     vAR_answers = []
#     vAR_final_sentences = []
#     for k,v in keyword_sentence_mapping.items():
#         if len(v)>0:
#             vAR_match = v[0].lower()
#             vAR_answers.append(k)
#             if k in vAR_match:
#                 vAR_temp = re.compile(re.escape(k), re.IGNORECASE)
#                 vAR_final_sentences.append(vAR_temp.sub('<answer>',vAR_match))
#             else:
#                 vAR_final_sentences.append(vAR_match)
#     return vAR_final_sentences, vAR_answers

# def printmd(string):
#     display(Markdown(string))

@st.cache(show_spinner=False)
@st.cache(allow_output_mutation=True)
def question(keyword_sentence_mapping):
    # tab = PrettyTable()
    vAR_final_sentences, vAR_answers  = sentence_answers(keyword_sentence_mapping)
    random.shuffle(vAR_answers)
    random.shuffle(vAR_final_sentences)
    vAR_cols_dict = {
        "A": vAR_answers,
        "B": vAR_final_sentences
    }
    pd.set_option("display.max_colwidth", None)
    vAR_cols = pd.DataFrame(vAR_cols_dict)
    return vAR_cols
    # tab.field_names=['A', 'B']
    # tab.align["A"] = "l"
    # tab.align["B"] = "l"

    # # printmd('**Match column A with column B**')
    # for word,context in zip(answers,final_sentences):
    #     tab.add_row([word,context.replace("\n"," ")])
    #     tab.add_row(['',''])
#     return tab
